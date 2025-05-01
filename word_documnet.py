from docx import Document
from docx.shared import Inches

# Create a new Word document
doc = Document()
doc.add_heading('Gaussian Beam Analysis from Position and Power Data', 0)

# Add introduction section
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "This report presents the analysis of experimental data consisting of power measurements "
    "as a function of angular position. The primary objective is to model the spatial distribution "
    "of the beam intensity using a Gaussian profile, determine the beam width (Full Width at Half Maximum, FWHM), "
    "and estimate the beam's wavelength under the assumption of a simplified Gaussian beam propagation model."
)

# Add data normalization section
doc.add_heading('Data Normalization', level=1)
doc.add_paragraph(
    "The raw position (in degrees) and power (in microwatts) measurements are normalized to values "
    "between 0 and 1. This is done to eliminate unit dependencies and ensure compatibility with the "
    "Gaussian fitting algorithm."
)

# Add method and fitting section
doc.add_heading('Gaussian Fitting', level=1)
doc.add_paragraph(
    "A Gaussian function of the form A * exp(-0.5 * ((x - mu) / sigma)^2) is fitted to the normalized data "
    "using non-linear least squares optimization. The fit provides estimates of amplitude (A), center (mu), "
    "and standard deviation (sigma). From the standard deviation, the Full Width at Half Maximum (FWHM) is "
    "calculated using the relation: FWHM = 2.355 * sigma."
)

# Add beam waist and wavelength estimation
doc.add_heading('Beam Waist and Wavelength Estimation', level=1)
doc.add_paragraph(
    "Assuming the fitted Gaussian profile corresponds to a cross-section of a Gaussian laser beam, "
    "the beam waist (w₀) is derived as w₀ = FWHM / 2.355. Under the assumption that the Rayleigh range (z_R) is known "
    "or approximated (e.g., z_R = 1 in normalized units), the wavelength λ can be estimated using the Gaussian beam "
    "relation: λ = (w₀² * π) / z_R."
)

# Add analysis and application
doc.add_heading('Analysis and Application', level=1)
doc.add_paragraph(
    "The Gaussian fit and the computed FWHM provide insight into the beam divergence and focus characteristics. "
    "In laboratory settings, this approach is valuable for characterizing optical systems, laser alignment, "
    "and determining beam quality. Deviations from a Gaussian profile may indicate aberrations or misalignment."
)

# Add conclusion section
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "This analysis offers a systematic method to extract Gaussian beam parameters from experimental data. "
    "The resulting fit parameters, particularly the FWHM and estimated wavelength, are critical for understanding "
    "the spatial and optical properties of the beam under investigation."
)

# Save the document
file_path = "Gaussian_Beam_Analysis_Report.docx"
doc.save(file_path)

file_path
